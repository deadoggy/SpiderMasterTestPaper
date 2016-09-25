#coding:utf-8
from UrlController import UrlController as UC
from docx import Document
from docx.enum.text import *
import urllib2
import urllib
from docx.shared import Inches
from docx.shared import RGBColor
import re

class SheetController:
    urlcontroller = UC()
    __sheetdic = urlcontroller.getAllSheetUrlBySubject()
    imageurl = {
        'math': 'http://kaoyan.eol.cn/shiti/shuxue/201605/',
        'english': 'http://kaoyan.eol.cn/shiti/yingyu/201605/',
        'politics': 'http://kaoyan.eol.cn/shiti/zhengzhi/201605/'
    }
    def __processSheet(self, url, subject):
        '''
        处理每一套试题
        :return:
        '''

        response = urllib2.urlopen(url + '.shtml')
        html = response.read().decode('gbk')
        pagecount = int(re.findall('''var _PAGE_COUNT="(.*?)"''', html)[0])
        title = re.findall(u'''<title>(.*?)-考研-中国教育在线</title>''', html)[0]
        docx = Document()

        for index in range(0, pagecount):
            if 0 == index:
                pageurl = url + '.shtml'
            else:
                pageurl = url + '_' + str(index) + '.shtml'
            self.__processPerSheet(pageurl, docx)

        docx.save(subject + '/' + title + '.docx')

    def __processPerSheet(self, url, docx):
        '''
        处理每一页的内容
        :param url: 该页的url
        :return:
        '''

        html = urllib2.urlopen(url).read().decode('gbk')

        pattern_1 = u'''<div class=TRS_Editor>(.*?)</div>'''
        pattern_1_2 =  u'''<div class="TRS_Editor">(.*?)</div>'''#网站上class不一致
        pattern_2 = u'''<p align="(.*?)">(.*?)</p>'''

        editor = re.findall(pattern_1, html, re.S)
        if [] == editor:
            editor = re.findall(pattern_1_2, html, re.S)
        contentset = re.findall(pattern_2, editor[0], re.S)

        for tag_tuple in contentset:
            tag = []
            tag.append(tag_tuple[0])
            tag.append(tag_tuple[1])
            para = docx.add_paragraph()
            if 'center' == tag[1]:
                para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            if [] != re.findall(u'&nbsp;', tag[1], re.S):#如果有空格
                tag[1] = re.sub(u'&nbsp;', u' ', tag[1] )

            if [] != re.findall(u'<strong>', tag[1]):#如果有黑体
                content = re.findall(u'<strong>(.*?)</strong>', tag[1], re.S)[0]
                #如果有<font>标签就设置颜色
                fontcontent = re.findall('''<font color="#ff0000">(.*?)</font>''', content, re.S)

                if [] != fontcontent:
                    run = para.add_run(fontcontent[0])
                    run.font.color.rgb = RGBColor(0x42, 0x24 , 0xE9)
                else:
                    run = docx.add_paragraph().add_run(content)
                #设置字体，加粗和下划线（？）
                run.font.italic = True
                run.font.bold = True
                run.font.underline = True

            elif [] != re.findall('<img', tag[1]):#如果是图片：
                #判断是什么学科的图片
                if [] != re.findall('zhengzhi', url):
                    imgroot = self.imageurl['politics']
                elif [] != re.findall('yingyu', url):
                    imgroot = self.imageurl['english']
                elif [] != re.findall('shuxue', url):
                    imgroot = self.imageurl['math']

                imgname = re.findall(''' src="./(.*?)" ''', tag[1])[0]
                imgurl = imgroot + imgname#获取图片链接
                urllib.urlretrieve(imgurl, 'image/' + imgname)
                docx.add_picture('image/' + imgname, width=Inches(5.96))
            else:
                para.add_run(tag[1])

        pass

    def getAllPaper(self):
        for key in self.__sheetdic.keys():
            if 'politics' == key or 'math' == key:
                continue
            for item in self.__sheetdic[key]:
                self.__processSheet(item, key)
        print "Done!"








